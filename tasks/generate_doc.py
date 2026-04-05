"""Generate HealthFlexx Platform Architecture document as .docx"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('HealthFlexx Coach — Platform Architecture', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Technical Reference Document\n'
    'HealthFlexx, LLC\n'
    'April 2026',
    style='Normal'
).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Platform Overview',
    '2. System Architecture',
    '3. Database Schema — All Tables',
    '4. Entity Relationship Diagram',
    '5. Feature Descriptions',
    '   5.1 User Onboarding',
    '   5.2 FIT Coaching Engine (Activity & Nutrition)',
    '   5.3 CBT-I Sleep Coaching Program',
    '   5.4 Real-Time Signal Analysis',
    '   5.5 Multi-Candidate Response Selection',
    '   5.6 Cross-Session Memory & Learning',
    '   5.7 Proactive Nudge System',
    '   5.8 In-App Messaging & Realtime Delivery',
    '   5.9 Background Notifications',
    '   5.10 Wellness Disclaimer & Consent',
    '   5.11 Health Data Integration',
    '   5.12 Session Insights & Strategy Learning',
    '6. Technology Stack',
    '7. Data Flow Diagrams',
    '8. Future Development Notes',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number' if not item.startswith('   ') else 'List Bullet')

doc.add_page_break()

# ============================================================
# 1. PLATFORM OVERVIEW
# ============================================================
doc.add_heading('1. Platform Overview', level=1)
doc.add_paragraph(
    'HealthFlexx Coach is an AI-powered wellness coaching platform that helps users improve '
    'their physical activity (steps), sleep habits, and nutrition through evidence-based behavioral '
    'techniques. The platform uses two primary coaching models:'
)
doc.add_paragraph('Functional Imagery Training (FIT) — for activity and nutrition coaching. '
    'Uses multi-sensory imagery, motivational interviewing signals, and commitment elicitation '
    'to drive behavior change.', style='List Bullet')
doc.add_paragraph('Cognitive Behavioral Therapy for Insomnia (CBT-I) — for sleep coaching. '
    'A structured 6-8 week program with sleep restriction therapy, stimulus control, cognitive '
    'restructuring, relaxation training, and sleep hygiene education.', style='List Bullet')
doc.add_paragraph(
    'The platform is not a medical device and does not provide medical advice. Users must '
    'accept a wellness coaching disclaimer before accessing the service. All health data '
    'is handled in accordance with HIPAA requirements.'
)

# ============================================================
# 2. SYSTEM ARCHITECTURE
# ============================================================
doc.add_heading('2. System Architecture', level=1)
doc.add_paragraph(
    'The platform consists of three main components:'
)

# Architecture table
arch_table = doc.add_table(rows=4, cols=3)
arch_table.style = 'Light Grid Accent 1'
arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Component', 'Technology', 'Hosting']
for i, h in enumerate(headers):
    arch_table.rows[0].cells[i].text = h
    arch_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data = [
    ['Backend API', 'Python 3.11 / FastAPI / Supabase / Claude API', 'Render'],
    ['Mobile App', 'Flutter / Dart / Supabase Flutter', 'Android (APK)'],
    ['Database', 'PostgreSQL via Supabase (RLS, Realtime, Auth)', 'Supabase Cloud'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        arch_table.rows[row_idx].cells[col_idx].text = val

doc.add_paragraph('')

# ============================================================
# 3. DATABASE SCHEMA
# ============================================================
doc.add_heading('3. Database Schema — All Tables', level=1)
doc.add_paragraph(
    'The platform uses 17 active tables in Supabase PostgreSQL. One table (messages) has been '
    'deprecated and dropped, with its functionality consolidated into chat_messages.'
)

# Tables table
tables_data = [
    ['#', 'Table', 'Migration', 'Purpose'],
    ['1', 'person', '002', 'Central identity — links to auth.users via Supabase Auth'],
    ['2', 'bot_profiles', '001', 'Per-user coaching config: domain, goals, baselines, gamification, learned preferences, timezone, disclaimer consent'],
    ['3', 'sessions', '001+003+008', 'Individual coaching sessions with phase tracking, outcomes, survey scores, conversation link'],
    ['4', 'conversations', '003', 'Long-lived chat threads (one per user per channel: in_app or sms)'],
    ['5', 'chat_messages', '003', 'All messages with delivery status tracking (sent → delivered → read), timestamps, session link'],
    ['6', 'commitments', '001', 'User commitments extracted from coaching sessions with follow-through tracking'],
    ['7', 'signal_logs', '001', 'Per-message signal snapshots: change_talk, sentiment, engagement, self_efficacy, imagery_engagement'],
    ['8', 'daily_progress', '001', 'Daily health metrics from wearable: steps, sleep, nutrition scores'],
    ['9', 'conversation_summaries', '004', 'Claude-generated conversation summaries with title, summary text, key topics'],
    ['10', 'session_insights', '005', 'FIT coaching insights: imagery practiced, cue-links, obstacles, strategy effectiveness'],
    ['11', 'nudge_log', '007', 'Nudge deduplication tracking — prevents sending same nudge type within 20-hour window'],
    ['12', 'sleep_programs', '008', 'CBT-I program state: phase, week number, sleep window, contraindications, modified flag'],
    ['13', 'sleep_diary_entries', '008', 'Nightly sleep diary: SOL, WASO, TST, TIB, SE%, quality rating, band data fusion'],
    ['14', 'isi_assessments', '008', 'Insomnia Severity Index scores (0-28) at screening, mid-program, end-of-program'],
    ['15', 'titration_log', '008', 'Sleep window adjustment history: weekly SE%, PTIB changes, bedtime recalculations'],
    ['16', 'sleep_screening_responses', '008', 'Safety screening results (PHI): contraindications detected, eligibility status'],
    ['17', 'messages (DROPPED)', '001', 'DEPRECATED — consolidated into chat_messages. No longer exists.'],
]

t = doc.add_table(rows=len(tables_data), cols=4)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for row_idx, row_data in enumerate(tables_data):
    for col_idx, val in enumerate(row_data):
        cell = t.rows[row_idx].cells[col_idx]
        cell.text = val
        if row_idx == 0:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph('')
doc.add_paragraph('RPC Functions:', style='Heading 3')
doc.add_paragraph('step_summary(user_id, interval_days) — Aggregates daily step data from pedometer table', style='List Bullet')
doc.add_paragraph('sleep_summary(user_id, interval_days) — Aggregates nightly sleep data', style='List Bullet')
doc.add_paragraph('get_last_band_sync(user_id) — Returns timestamp of most recent wearable data push', style='List Bullet')
doc.add_paragraph('sleep_diary_weekly_summary(person_id, week_start) — Weekly diary aggregates for titration decisions', style='List Bullet')
doc.add_paragraph('sleep_diary_trend(person_id) — Weekly SE% trend over program duration', style='List Bullet')

# ============================================================
# 4. ENTITY RELATIONSHIP DIAGRAM
# ============================================================
doc.add_heading('4. Entity Relationship Diagram', level=1)

erd = """
                            ┌─────────────┐
                            │ auth.users  │
                            │─────────────│
                            │ id (PK)     │
                            │ email       │
                            └──────┬──────┘
                                   │ 1:1
                                   ▼
                            ┌─────────────┐
                            │   person    │
                            │─────────────│
                            │ id (PK/FK)  │
                            │ first_name  │
                            │ email       │
                            │ phone       │
                            └──────┬──────┘
                                   │
           ┌───────────┬───────────┼───────────┬──────────────┐
           │           │           │           │              │
           ▼ 1:1       ▼ 1:N      ▼ 1:N      ▼ 1:N          ▼ 1:1
   ┌──────────────┐ ┌────────┐ ┌──────────┐ ┌───────────┐ ┌──────────────┐
   │ bot_profiles │ │sessions│ │commitments│ │conversations│ │sleep_programs│
   │──────────────│ │────────│ │──────────│ │───────────│ │──────────────│
   │ person_id PK │ │ id PK  │ │ id PK    │ │ id PK     │ │ id PK        │
   │ domain       │ │person ▲│ │person ▲  │ │person ▲   │ │ person_id FK │
   │ goal         │ │convo ▲ │ │session ▲ │ │ channel   │ │ program_phase│
   │ is_onboarded │ │ domain │ │ descrip  │ │           │ │ week_number  │
   │ timezone     │ │ phase  │ │ followed │ │           │ │ is_modified  │
   │ sleep_prog ▲ │ │ survey │ │ confid   │ │           │ │ ptib_minutes │
   │ disclaimer_at│ │prog_ph │ └──────────┘ │           │ │ bedtime      │
   │ effective_*  │ └───┬────┘              │           │ │ wake_time    │
   └──────────────┘     │                   │           │ └──────┬───────┘
                        │                   │           │        │
              ┌─────────┼──────────┐        │           │   ┌────┴─────────┐
              │         │          │        ▼ 1:N       │   │              │
              ▼ 1:N     ▼ 1:1     ▼ 1:N  ┌───────────┐ │   ▼ 1:N         ▼ 1:N
   ┌────────────┐ ┌──────────┐ ┌──────┐  │chat_msgs  │ │ ┌────────────┐ ┌──────────┐
   │signal_logs │ │ session  │ │nudge │  │───────────│ │ │sleep_diary │ │titration │
   │────────────│ │ insights │ │ _log │  │ id PK     │ │ │_entries    │ │  _log    │
   │ id PK      │ │──────────│ │──────│  │convo_id ▲ │ │ │────────────│ │──────────│
   │ session ▲  │ │ id PK    │ │id PK │  │session ▲  │ │ │ id PK      │ │ id PK    │
   │ change_talk│ │session ▲ │ │pers▲ │  │sender_type│ │ │ person ▲   │ │ person ▲ │
   │ sentiment  │ │person ▲  │ │ type │  │ content   │ │ │ diary_date │ │ week_num │
   │ engagement │ │ imagery  │ │msg ▲ │  │ status    │ │ │ sol, waso  │ │ se_pct   │
   │ self_eff   │ │ cue_links│ └──────┘  │delivered_at│ │ │ tst, tib   │ │ ptib_old │
   │ imagery_eng│ │ obstacles│           │ read_at   │ │ │ se_pct     │ │ ptib_new │
   │ recomm_act │ │ strat_*  │           └───────────┘ │ │ quality    │ │ adjust   │
   └────────────┘ └──────────┘                         │ └────────────┘ └──────────┘
                                                       │
                        ┌──────────────────────────────┘
                        │                         ┌──────────────────┐
                        ▼ 1:1                     │                  │
              ┌──────────────────┐                ▼ 1:N              ▼ 1:1
              │ conversation     │         ┌──────────────┐  ┌───────────────┐
              │ _summaries       │         │isi_assessments│  │sleep_screening│
              │──────────────────│         │──────────────│  │ _responses    │
              │ id PK            │         │ id PK        │  │───────────────│
              │ conversation ▲   │         │ person ▲     │  │ id PK         │
              │ title            │         │ session ▲    │  │ person_id FK  │
              │ summary          │         │ week_number  │  │ responses     │
              │ key_topics       │         │ item_scores  │  │ contras       │
              └──────────────────┘         │ total_score  │  │ eligible_*    │
                                           │ severity     │  └───────────────┘
              ┌──────────────────┐         └──────────────┘
              │ daily_progress   │
              │──────────────────│
              │ id PK            │
              │ person ▲         │
              │ date             │
              │ steps, sleep     │
              │ nutrition_score  │
              └──────────────────┘

LEGEND:  ▲ = FK reference    PK = Primary Key    1:N = one-to-many    1:1 = one-to-one
"""

p = doc.add_paragraph()
run = p.add_run(erd)
run.font.name = 'Consolas'
run.font.size = Pt(7)

# ============================================================
# 5. FEATURE DESCRIPTIONS
# ============================================================
doc.add_heading('5. Feature Descriptions', level=1)

# 5.1
doc.add_heading('5.1 User Onboarding', level=2)
doc.add_paragraph(
    'New users go through a conversational onboarding flow that gathers four data points: '
    'health goal, chosen domain (activity/nutrition/sleep), biggest barrier, and core values. '
    'Claude conducts this as a natural conversation, not a survey. Responses are extracted '
    'via classification calls and stored in bot_profiles.'
)
doc.add_paragraph(
    'Sleep domain users receive additional safety screening questions (bipolar, epilepsy, '
    'sleep apnea, fall risk). If contraindications are detected, the user is enrolled in a '
    'modified CBT-I program that excludes sleep restriction therapy. Screening responses are '
    'stored in sleep_screening_responses.'
)
doc.add_paragraph('Key files: bot/onboarding.py, bot/engine.py (_handle_onboarding, _setup_sleep_program)')

# 5.2
doc.add_heading('5.2 FIT Coaching Engine (Activity & Nutrition)', level=2)
doc.add_paragraph(
    'For activity and nutrition domains, the bot uses Functional Imagery Training (FIT) — '
    'an evidence-based technique that uses multi-sensory mental imagery to build and sustain '
    'motivation for behavior change.'
)
doc.add_paragraph('Session Phases (7):', style='Heading 4')
phases = [
    ('check_in', 'Brief assessment of current state and last commitment follow-through (2 messages)'),
    ('celebrate', 'Recognize effort, build identity as "someone who follows through" (2 messages)'),
    ('education', 'Share relevant health insight matched to domain and data (1 message)'),
    ('fit_imagery', 'Guided multi-sensory imagery — the core technique. Signal-gated: advances when change_talk > 0.6 or after 8 messages'),
    ('commitment', 'Elicit specific, confidence-rated commitment with implementation details'),
    ('gamification', 'Brief momentum checkpoint or streak celebration (1 message)'),
    ('close', 'Affirm progress, set up next session touchpoint (2 messages)'),
]
for phase, desc in phases:
    doc.add_paragraph(f'{phase}: {desc}', style='List Bullet')

doc.add_paragraph('Techniques (5):', style='Heading 4')
techniques = [
    ('fit_imagery', 'Multi-sensory process/outcome imagery, channel switching, Choice Point rehearsal'),
    ('commitment', 'Implementation intention elicitation, confidence scaling, habit stacking'),
    ('reframing', 'Question assumptions gently, offer alternative perspective, connect to values'),
    ('behavioral_economics', 'Nudge framing, anchoring, loss aversion, immediate reward salience'),
    ('identity', '"Becoming someone who..." language, reflect actions as identity evidence'),
]
for tech, desc in techniques:
    doc.add_paragraph(f'{tech}: {desc}', style='List Bullet')

doc.add_paragraph('Key files: bot/engine.py, bot/session.py, bot/prompts/system.py, bot/prompts/fit_coaching_prompt.py, strategies/adaptation.py')

# 5.3
doc.add_heading('5.3 CBT-I Sleep Coaching Program', level=2)
doc.add_paragraph(
    'For sleep domain users, the platform delivers a structured 6-8 week Cognitive Behavioral '
    'Therapy for Insomnia (CBT-I) program — the AASM-recommended first-line treatment for '
    'chronic insomnia with 70-80% response rates.'
)

doc.add_paragraph('Program Phases:', style='Heading 4')
prog_phases = [
    ('screening', 'Week 0: ISI assessment + safety screening during onboarding'),
    ('baseline', 'Week 1: Collect 7+ nights of sleep diary data'),
    ('sleep_restriction', 'Week 2: 3P model education, calculate sleep window, introduce stimulus control'),
    ('titration_hygiene', 'Week 3: First titration + sleep hygiene education'),
    ('titration_relax', 'Week 4: Titration + progressive muscle relaxation'),
    ('titration_cog1', 'Week 5: Titration + cognitive restructuring (identify beliefs)'),
    ('titration_cog2', 'Week 6: Titration + cognitive restructuring (Socratic questioning)'),
    ('integration', 'Week 7-8: Relapse prevention, self-management, graduation'),
]
for phase, desc in prog_phases:
    doc.add_paragraph(f'{phase}: {desc}', style='List Bullet')

doc.add_paragraph('Modified Program (contraindicated users — no sleep restriction):', style='Heading 4')
doc.add_paragraph(
    'screening → baseline → stimulus_control → relaxation → cognitive1 → cognitive2 → integration. '
    'Skips all sleep restriction and titration. Focuses on stimulus control, cognitive restructuring, '
    'relaxation training, and sleep hygiene.'
)

doc.add_paragraph('Three Session Types:', style='Heading 4')
doc.add_paragraph('Morning diary (~2 min): Conversational sleep data collection. Bot asks about bedtime, time to fall asleep, awakenings, wake time, quality. Extracts structured data via Claude. Calculates TIB, TST, SE%.', style='List Bullet')
doc.add_paragraph('Weekly module (~15 min): Structured session with phases: check_in → diary_review → titration → education → practice → planning → close. Delivers week-specific CBT-I content.', style='List Bullet')
doc.add_paragraph('On-demand: User-initiated. Responsive coaching drawing from whichever CBT-I technique is most relevant.', style='List Bullet')

doc.add_paragraph('Sleep Window Algorithm:', style='Heading 4')
doc.add_paragraph('Initial: PTIB = average TST from 7+ baseline nights. Floor: 4.5 hours (270 min). User picks fixed wake time. Bedtime = wake_time - PTIB.', style='List Bullet')
doc.add_paragraph('Weekly titration: SE% > 90% → expand 15 min. SE% 85-90% → hold. SE% < 85% → compress 15 min.', style='List Bullet')
doc.add_paragraph('Stop when desired TST reached with SE% >= 85%.', style='List Bullet')

doc.add_paragraph('CBT-I Techniques (6):', style='Heading 4')
cbti_techs = [
    ('sleep_restriction', 'Compress time in bed to match actual sleep, titrate weekly based on SE%'),
    ('stimulus_control', '6 rules: bed=sleep only, get up if awake 20 min, fixed wake time, no naps'),
    ('cognitive_restructuring', 'Socratic questioning on catastrophizing, effort beliefs, sleep expectations'),
    ('relaxation_training', 'PMR (16→7→4 muscle groups), 4-7-8 breathing, body scan'),
    ('sleep_hygiene', 'Caffeine, alcohol, temperature, light, exercise timing, nap rules'),
    ('psychoeducation', '3P model, sleep homeostasis, circadian rhythm, normal sleep architecture'),
]
for tech, desc in cbti_techs:
    doc.add_paragraph(f'{tech}: {desc}', style='List Bullet')

doc.add_paragraph('Key files: bot/sleep_handler.py, bot/sleep_program.py, bot/sleep_diary.py, bot/prompts/sleep_cbti.py, signals/sleep_classifier.py')

# 5.4
doc.add_heading('5.4 Real-Time Signal Analysis', level=2)
doc.add_paragraph(
    'Every user message is analyzed in real-time by 4 Claude-powered classifiers that produce '
    'a SignalSnapshot. These signals drive technique selection and phase advancement.'
)

signals_table_data = [
    ['Signal', 'Range', 'Classifier', 'What It Detects'],
    ['change_talk_ratio', '0-1', 'Claude (DARN-C)', 'Language arguing FOR behavior change'],
    ['sustain_talk_ratio', '0-1', 'Claude (DARN-C)', 'Language arguing AGAINST change (resistance)'],
    ['sentiment', '-1 to 1', 'Claude', 'Emotional valence, energy, openness, hope'],
    ['sentiment_delta', '-2 to 2', 'Computed', 'Change from last message — key signal for technique switching'],
    ['engagement', '0-1', 'Claude', 'Response depth, personal detail, questions asked'],
    ['commitment_strength', '0-1', 'Claude', 'Strength of commitment language ("I will" vs "maybe")'],
    ['self_efficacy', '0-1', 'Claude', 'Confidence in ability to make changes'],
    ['imagery_engagement', '0-1', 'Claude', 'Sensory detail and vividness in imagery responses'],
]

st = doc.add_table(rows=len(signals_table_data), cols=4)
st.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(signals_table_data):
    for col_idx, val in enumerate(row_data):
        cell = st.rows[row_idx].cells[col_idx]
        cell.text = val
        if row_idx == 0:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph('')
doc.add_paragraph('Sleep-Specific Signals (6 additional):', style='Heading 4')
sleep_signals = [
    ('sleep_anxiety', 'Catastrophizing about sleep consequences'),
    ('effort_to_sleep', '"Trying harder" to sleep (paradoxically prevents it)'),
    ('bed_frustration', 'Frustration about being awake in bed'),
    ('safety_concern', 'Mentions of contraindicated conditions or crisis'),
    ('adherence_difficulty', 'Struggling with sleep restriction / stimulus control protocol'),
    ('relaxation_engagement', 'Participating in guided relaxation exercise'),
]
for sig, desc in sleep_signals:
    doc.add_paragraph(f'{sig}: {desc}', style='List Bullet')

doc.add_paragraph('Key files: signals/analyzer.py, signals/change_talk.py, signals/sentiment.py, signals/engagement.py, signals/commitment.py, signals/sleep_classifier.py')

# 5.5
doc.add_heading('5.5 Multi-Candidate Response Selection', level=2)
doc.add_paragraph(
    'The bot generates N response candidates (default 3) with stylistic variation, scores each '
    'against the current signal state (change talk, sentiment, engagement, values alignment), '
    'and returns the highest-scoring one. This is inspired by the Reddit ChangeMyView study '
    'where tournament-style selection significantly improved persuasion quality.'
)
doc.add_paragraph('Key files: integrations/claude.py (generate_candidates, score_candidate, select_best_response)')

# 5.6
doc.add_heading('5.6 Cross-Session Memory & Learning', level=2)
doc.add_paragraph(
    'The platform maintains coaching memory across sessions to provide continuity:'
)
doc.add_paragraph('Recent session insights: imagery practiced, cue-links established, obstacles identified', style='List Bullet')
doc.add_paragraph('Commitment track record: follow-through percentage, patterns', style='List Bullet')
doc.add_paragraph('Effective vs. ineffective strategies: learned from signal_logs sentiment deltas', style='List Bullet')
doc.add_paragraph('Preferred imagery style: detected and stored for personalization', style='List Bullet')
doc.add_paragraph('Sleep program state: current phase, sleep window, diary trends, ISI trajectory', style='List Bullet')
doc.add_paragraph(
    'After each session closes, a background task extracts insights via Claude and updates '
    'learned_preferences in bot_profiles. This data is injected into the system prompt for '
    'the next session.'
)
doc.add_paragraph('Key files: services/insights_extractor.py, bot/prompts/memory.py, db/client.py (load_coaching_memory, load_sleep_coaching_memory)')

# 5.7
doc.add_heading('5.7 Proactive Nudge System', level=2)
doc.add_paragraph(
    'APScheduler runs inside FastAPI\'s event loop, evaluating nudge rules every 30 minutes '
    'for all active users. Max 1 nudge per user per cycle. All nudges have a 20-hour cooldown '
    'via nudge_log deduplication. Messages are timezone-aware.'
)

nudges_data = [
    ['Nudge', 'Domain', 'Window (local)', 'Trigger'],
    ['morning_sync_reminder', 'All', '7-9 AM', 'Band not synced today'],
    ['evening_sync_reminder', 'All', '7-9 PM', 'Band not synced since noon'],
    ['steps_nudge', 'Activity', '6:30-7:30 PM', '300+ steps behind daily avg'],
    ['cue_link_reminder', 'Activity/Nutrition', '7-8 AM', 'Active cue-link, no recent message'],
    ['commitment_checkin', 'Activity/Nutrition', '10 AM-12 PM', 'Yesterday commitment, no follow-through'],
    ['milestone_celebration', 'Activity', '9 AM-8 PM', 'Steps crossed 5k/7.5k/10k'],
    ['morning_diary_reminder', 'Sleep', '7-9 AM', 'No diary entry for last night'],
    ['wind_down_reminder', 'Sleep', '1hr before bedtime', 'Active sleep program'],
    ['bedtime_reminder', 'Sleep', '15min before bedtime', 'Active sleep program'],
    ['weekly_session_reminder', 'Sleep', '10 AM-12 PM', 'Module not completed this week'],
]

nt = doc.add_table(rows=len(nudges_data), cols=4)
nt.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(nudges_data):
    for col_idx, val in enumerate(row_data):
        cell = nt.rows[row_idx].cells[col_idx]
        cell.text = val
        if row_idx == 0:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph('')
doc.add_paragraph('Key files: services/nudge_scheduler.py')

# 5.8
doc.add_heading('5.8 In-App Messaging & Realtime Delivery', level=2)
doc.add_paragraph(
    'Two delivery mechanisms work together:'
)
doc.add_paragraph('Supabase Realtime (foreground): Instant message delivery when app is open. Subscribes to chat_messages INSERT events filtered by conversation_id and sender_type=bot. Deduplicates by message ID.', style='List Bullet')
doc.add_paragraph('Background polling (closed/backgrounded): WorkManager polls /api/v1/unread-count every 15 minutes. Shows local notification if unread messages exist with status "sent".', style='List Bullet')
doc.add_paragraph(
    'Message lifecycle: sent (backend creates) → delivered (app receives, sets delivered_at) → '
    'read (user views chat screen, sets read_at).'
)
doc.add_paragraph('Key files: providers/chat_provider.dart, services/chat_service.dart, services/background_service.dart')

# 5.9
doc.add_heading('5.9 Background Notifications', level=2)
doc.add_paragraph(
    'flutter_local_notifications handles local notification display on Android. Notifications use '
    'the "coach_nudges" channel with high importance/priority (sound + vibration). The app requests '
    'POST_NOTIFICATIONS permission on first launch. Tapping a notification opens the app and '
    'refreshes the chat.'
)
doc.add_paragraph('Key files: services/notification_service.dart, services/background_service.dart')

# 5.10
doc.add_heading('5.10 Wellness Disclaimer & Consent', level=2)
doc.add_paragraph(
    'A consent screen appears after first login and before accessing the app. Users must tap '
    '"I Agree" to proceed. The disclaimer covers: not medical advice, not a healthcare provider, '
    'consult your doctor, no guarantee of results, emergency situations, user responsibility, '
    'data privacy, and informed consent. The acceptance timestamp is stored in '
    'bot_profiles.disclaimer_accepted_at. The screen does not appear again once accepted.'
)
doc.add_paragraph('Key files: screens/consent_screen.dart, app.dart (_ConsentGate)')

# 5.11
doc.add_heading('5.11 Health Data Integration', level=2)
doc.add_paragraph(
    'Health data flows from external wearable devices through Supabase RPC functions:'
)
doc.add_paragraph('Step data: step_summary RPC returns 7-day average and daily breakdown from pedometer table', style='List Bullet')
doc.add_paragraph('Sleep data: sleep_summary RPC returns nightly duration and quality from sleep_stats table', style='List Bullet')
doc.add_paragraph('Band sync: get_last_band_sync RPC returns most recent data push timestamp', style='List Bullet')
doc.add_paragraph('Sleep diary: Conversational collection fused with band data. Diary values take priority for SOL and bedtime (subjective experience matters for CBT-I). Band fills gaps.', style='List Bullet')
doc.add_paragraph(
    'Health data is injected into the system prompt so Claude can reference specific numbers '
    'naturally in coaching conversations.'
)
doc.add_paragraph('Key files: db/client.py (get_health_summary, get_step_summary, get_sleep_summary, get_last_band_sync)')

# 5.12
doc.add_heading('5.12 Session Insights & Strategy Learning', level=2)
doc.add_paragraph(
    'After each session closes, background tasks run:'
)
doc.add_paragraph('Insight extraction: Claude analyzes the full conversation + signal trajectory to extract imagery practiced, cue-links, obstacles, and strategy effectiveness. Saved to session_insights.', style='List Bullet')
doc.add_paragraph('Strategy scoring: Signal deltas after each technique application determine what was effective vs. ineffective. Updates learned_preferences in bot_profiles.', style='List Bullet')
doc.add_paragraph('Session scoring: Claude rates the session on motivation (1-5), surprise (text), and helpfulness (1-5). Saved to sessions survey columns.', style='List Bullet')
doc.add_paragraph('Commitment extraction: During commitment phase, Claude detects if a commitment was made and extracts description, confidence, and timing. Saved to commitments table.', style='List Bullet')
doc.add_paragraph('Conversation summary: Claude generates a title, summary, and key topics. Saved to conversation_summaries.', style='List Bullet')
doc.add_paragraph('Key files: services/insights_extractor.py, services/summarizer.py, bot/engine.py (_extract_and_learn, _score_session, _try_extract_commitment)')

# ============================================================
# 6. TECHNOLOGY STACK
# ============================================================
doc.add_heading('6. Technology Stack', level=1)

stack_data = [
    ['Layer', 'Technology', 'Version/Detail'],
    ['Backend Framework', 'FastAPI', 'Python 3.11+'],
    ['AI/LLM', 'Anthropic Claude API', 'claude-3-5-sonnet via anthropic SDK'],
    ['Database', 'PostgreSQL via Supabase', 'With RLS, Realtime, Auth'],
    ['Mobile Framework', 'Flutter / Dart', 'SDK 3.11+'],
    ['State Management', 'Provider', 'ChangeNotifier pattern'],
    ['Auth', 'Supabase Auth', 'Email/password, JWT'],
    ['Background Jobs', 'APScheduler (backend)', 'AsyncIOScheduler, 30-min interval'],
    ['Background Polling', 'WorkManager (Flutter)', '15-min periodic task'],
    ['Notifications', 'flutter_local_notifications', 'Android local notifications'],
    ['Realtime', 'Supabase Realtime', 'PostgreSQL CDC on chat_messages'],
    ['Hosting', 'Render', 'Docker container, auto-deploy from GitHub'],
    ['Scheduling', 'APScheduler', 'In-process, no external scheduler needed'],
]

stack_t = doc.add_table(rows=len(stack_data), cols=3)
stack_t.style = 'Light Grid Accent 1'
for row_idx, row_data in enumerate(stack_data):
    for col_idx, val in enumerate(row_data):
        cell = stack_t.rows[row_idx].cells[col_idx]
        cell.text = val
        if row_idx == 0:
            cell.paragraphs[0].runs[0].bold = True

# ============================================================
# 7. DATA FLOW DIAGRAMS
# ============================================================
doc.add_heading('7. Data Flow Diagrams', level=1)

doc.add_paragraph('Chat Message Flow:', style='Heading 3')
chat_flow = """User sends message (Flutter)
    → POST /api/v1/chat (FastAPI)
        → save_chat_message(user msg, session_id) → chat_messages
        → engine.handle_message(person_id, message)
            → Signal analysis (4 Claude classifiers) → signal_logs
            → Phase advancement check
            → Strategy selection (signal-driven)
            → build_system_prompt (profile + phase + domain + health data + memory + signals)
            → Claude multi-candidate generation → score → best response
            → Commitment extraction (if commitment phase)
        → save_chat_message(bot response, session_id) → chat_messages
    ← Response to Flutter
        → Supabase Realtime delivers to other clients
        → Background poll picks up if app closed"""

p = doc.add_paragraph()
run = p.add_run(chat_flow)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph('')
doc.add_paragraph('Session Close Flow:', style='Heading 3')
close_flow = """Session reaches "close" phase with 2+ messages
    → _close_session()
        → Update session: ended_at, phase_reached, commitment_made, confidence_score
        → Background: _score_session()
            → Claude rates motivation/surprise/helpfulness → sessions
        → Background: _extract_and_learn()
            → Fetch session messages + signal_logs
            → InsightsExtractor.extract() → session_insights
            → Update learned_preferences → bot_profiles
            → ConversationSummarizer.summarize() → conversation_summaries"""

p = doc.add_paragraph()
run = p.add_run(close_flow)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# ============================================================
# 8. FUTURE DEVELOPMENT NOTES
# ============================================================
doc.add_heading('8. Future Development Notes', level=1)

doc.add_paragraph('Architecture Decisions to Preserve:', style='Heading 3')
doc.add_paragraph('Domain routing in engine.py: sleep domain → SleepSessionHandler, all others → FIT flow. New domains should follow this pattern.', style='List Bullet')
doc.add_paragraph('Signal-driven technique selection: Signals drive strategy, not hardcoded phase-technique mappings. New signals should integrate into the existing adaptation framework.', style='List Bullet')
doc.add_paragraph('Multi-candidate response selection: Quality over speed. The N-candidate tournament is a core differentiator.', style='List Bullet')
doc.add_paragraph('Background task pattern: Non-blocking work (insights, summaries, scoring) runs via asyncio.create_task after session close.', style='List Bullet')
doc.add_paragraph('Nudge deduplication: All nudges go through nudge_log with 20-hour cooldown. Never bypass this.', style='List Bullet')
doc.add_paragraph('Safety screening: Sleep restriction is gated on screening. Any new behavioral intervention with contraindications should follow the same screen-and-gate pattern.', style='List Bullet')

doc.add_paragraph('Known Limitations / Future Work:', style='Heading 3')
doc.add_paragraph('Push notifications: Currently uses local notifications via 15-min WorkManager polling. FCM (Firebase Cloud Messaging) would provide near-instant delivery.', style='List Bullet')
doc.add_paragraph('iOS support: App is Android-only. Flutter supports iOS but needs Apple Developer account, notification entitlements, and iOS-specific WorkManager alternative.', style='List Bullet')
doc.add_paragraph('Sleep diary automation: Currently conversational. Could integrate with Health Connect API to auto-fill some fields.', style='List Bullet')
doc.add_paragraph('ISI delivery: Defined in prompts but not yet wired into the session flow as a distinct phase. Needs integration into sleep_handler.py.', style='List Bullet')
doc.add_paragraph('Program phase advancement: Currently manual (time-based check). Could be automated via a weekly cron job that advances all eligible users.', style='List Bullet')
doc.add_paragraph('daily_progress table: Schema exists but is not actively populated. Could be filled by a daily aggregation job from pedometer/sleep data.', style='List Bullet')
doc.add_paragraph('Gamification: Points, streaks, and badges columns exist in bot_profiles but are not implemented in the coaching flow.', style='List Bullet')
doc.add_paragraph('Signal calibration: change_talk and sentiment may return binary (0/1) instead of gradients. Prompt tuning or few-shot examples may improve granularity.', style='List Bullet')
doc.add_paragraph('Nutrition coaching: Domain exists but has no nutrition-specific data integration (meal logging, nutrition scores). Currently uses generic FIT coaching.', style='List Bullet')
doc.add_paragraph('Multi-user testing: Platform has been tested with 3 users. Load testing and concurrent session handling need validation at scale.', style='List Bullet')

# Save
output_path = os.path.expanduser('~/behavior_change_app/HealthFlexx_Platform_Architecture.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
